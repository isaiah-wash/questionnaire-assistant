"""
Document Parser Module
Parses Excel, CSV, Word, and PDF files to extract Q&A pairs.
"""

import io
import re
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document
from pypdf import PdfReader
from anthropic import Anthropic


@dataclass
class QAPair:
    question: str
    answer: str
    source_file: str
    category: str = ""


class DocumentParser:
    def __init__(self, anthropic_api_key: str):
        self.client = Anthropic(api_key=anthropic_api_key)

    def parse_file(self, file_path: str, file_content: bytes = None, extract_questions_only: bool = False) -> list[QAPair]:
        """Parse a file and extract Q&A pairs.

        Args:
            file_path: Path to the file
            file_content: File content as bytes
            extract_questions_only: If True, extract questions even if answers are empty
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix in [".xlsx", ".xls"]:
            return self._parse_excel(file_path, file_content, extract_questions_only)
        elif suffix == ".csv":
            return self._parse_csv(file_path, file_content, extract_questions_only)
        elif suffix == ".docx":
            return self._parse_word(file_path, file_content, extract_questions_only)
        elif suffix == ".pdf":
            return self._parse_pdf(file_path, file_content, extract_questions_only)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def _parse_excel(self, file_path: str, file_content: bytes = None, extract_questions_only: bool = False) -> list[QAPair]:
        """Parse Excel file to extract Q&A pairs."""
        if file_content:
            df = pd.read_excel(io.BytesIO(file_content), sheet_name=None)
        else:
            df = pd.read_excel(file_path, sheet_name=None)

        qa_pairs = []
        source_name = Path(file_path).name

        for sheet_name, sheet_df in df.items():
            pairs = self._extract_qa_from_dataframe(sheet_df, source_name, sheet_name, extract_questions_only)
            qa_pairs.extend(pairs)

        return qa_pairs

    def _parse_csv(self, file_path: str, file_content: bytes = None, extract_questions_only: bool = False) -> list[QAPair]:
        """Parse CSV file to extract Q&A pairs."""
        if file_content:
            df = pd.read_csv(io.BytesIO(file_content))
        else:
            df = pd.read_csv(file_path)

        source_name = Path(file_path).name
        return self._extract_qa_from_dataframe(df, source_name, "", extract_questions_only)

    def _extract_qa_from_dataframe(
        self, df: pd.DataFrame, source_name: str, category: str = "", extract_questions_only: bool = False
    ) -> list[QAPair]:
        """Extract Q&A pairs from a DataFrame by detecting question/answer columns."""
        qa_pairs = []

        # Clean column names
        df.columns = df.columns.str.strip().str.lower()

        # Try to find question and answer columns
        question_col = None
        answer_col = None

        question_patterns = ["question", "query", "ask", "q", "requirement", "item"]
        answer_patterns = ["answer", "response", "reply", "a", "value", "input"]

        for col in df.columns:
            col_lower = str(col).lower()
            if question_col is None:
                for pattern in question_patterns:
                    if pattern in col_lower:
                        question_col = col
                        break
            if answer_col is None:
                for pattern in answer_patterns:
                    if pattern in col_lower:
                        answer_col = col
                        break

        # Extract questions (with or without answers)
        if question_col:
            for _, row in df.iterrows():
                q = str(row[question_col]).strip()

                # Skip empty or NaN questions
                if not q or q.lower() == "nan":
                    continue

                # Get answer if available
                a = ""
                if answer_col:
                    a = str(row[answer_col]).strip()
                    if a.lower() == "nan":
                        a = ""

                # If extracting questions only, include even without answers
                # If extracting Q&A pairs, require both
                if extract_questions_only or a:
                    qa_pairs.append(
                        QAPair(
                            question=q,
                            answer=a,
                            source_file=source_name,
                            category=category,
                        )
                    )
        else:
            # Use Claude to extract Q&A pairs from unstructured data
            text = df.to_string()
            qa_pairs = self._extract_qa_with_claude(text, source_name, category, extract_questions_only)

        return qa_pairs

    def _parse_word(self, file_path: str, file_content: bytes = None, extract_questions_only: bool = False) -> list[QAPair]:
        """Parse Word document to extract Q&A pairs."""
        if file_content:
            doc = Document(io.BytesIO(file_content))
        else:
            doc = Document(file_path)

        source_name = Path(file_path).name

        # Extract all text from paragraphs and tables
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)

        # Use Claude to extract Q&A pairs
        return self._extract_qa_with_claude(full_text, source_name, "", extract_questions_only)

    def _parse_pdf(self, file_path: str, file_content: bytes = None, extract_questions_only: bool = False) -> list[QAPair]:
        """Parse PDF file to extract Q&A pairs."""
        if file_content:
            reader = PdfReader(io.BytesIO(file_content))
        else:
            reader = PdfReader(file_path)

        source_name = Path(file_path).name

        # Extract all text
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        full_text = "\n".join(text_parts)

        # Use Claude to extract Q&A pairs
        return self._extract_qa_with_claude(full_text, source_name, "", extract_questions_only)

    def _extract_qa_with_claude(
        self, text: str, source_name: str, category: str = "", extract_questions_only: bool = False
    ) -> list[QAPair]:
        """Use Claude to extract Q&A pairs from unstructured text."""
        # Truncate if too long
        if len(text) > 50000:
            text = text[:50000] + "\n...[truncated]"

        if extract_questions_only:
            prompt = f"""Extract all questions from this due diligence/compliance questionnaire document.
This is a NEW questionnaire to be filled out, so answers may be missing or empty.

Return a JSON array of objects with "question" and "answer" fields.
For questions without answers, set "answer" to an empty string.
Only include actual questions, not headers or instructions.

Document text:
{text}

Return ONLY valid JSON array, no other text. Example format:
[{{"question": "What is your company name?", "answer": ""}}]"""
        else:
            prompt = f"""Extract all question-answer pairs from this due diligence/compliance questionnaire document.

Return a JSON array of objects with "question" and "answer" fields.
Only include actual Q&A pairs, not headers or instructions.
If a question has no answer, skip it.

Document text:
{text}

Return ONLY valid JSON array, no other text. Example format:
[{{"question": "What is your company name?", "answer": "Acme Corp"}}]"""

        response = self.client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}],
        )

        try:
            import json

            result_text = response.content[0].text.strip()

            # Try to extract JSON from response
            if result_text.startswith("["):
                pairs_data = json.loads(result_text)
            else:
                # Try to find JSON array in response
                match = re.search(r"\[.*\]", result_text, re.DOTALL)
                if match:
                    pairs_data = json.loads(match.group())
                else:
                    return []

            qa_pairs = []
            for item in pairs_data:
                if "question" in item and "answer" in item:
                    q = str(item["question"]).strip()
                    a = str(item["answer"]).strip()
                    if q and a:
                        qa_pairs.append(
                            QAPair(
                                question=q,
                                answer=a,
                                source_file=source_name,
                                category=category,
                            )
                        )

            return qa_pairs

        except (json.JSONDecodeError, IndexError, KeyError):
            return []
