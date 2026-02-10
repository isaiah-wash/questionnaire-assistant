"""
Exporter Module
Exports filled questionnaires back to their original format.
"""

import io
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import RGBColor
from openpyxl import load_workbook
from openpyxl.styles import PatternFill


class QuestionnaireExporter:
    def __init__(self):
        pass

    def export_to_excel(
        self,
        template_content: bytes,
        template_filename: str,
        filled_answers: list[dict],
    ) -> bytes:
        """
        Export filled answers to Excel format.

        Args:
            template_content: Original Excel file content
            template_filename: Original filename
            filled_answers: List of {question, answer, confidence, needs_review}

        Returns:
            Filled Excel file as bytes
        """
        # Load the original Excel file
        df_dict = pd.read_excel(io.BytesIO(template_content), sheet_name=None)

        # Create a mapping of questions to answers
        answer_map = {qa["question"]: qa for qa in filled_answers}

        output = io.BytesIO()

        # Write to Excel with openpyxl for styling
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            for sheet_name, df in df_dict.items():
                # Clean column names
                df.columns = df.columns.str.strip()

                # Find question and answer columns
                question_col = None
                answer_col = None

                question_patterns = ["question", "query", "ask", "q", "requirement", "item"]
                answer_patterns = ["answer", "response", "reply", "a", "value", "input"]

                for col in df.columns:
                    col_lower = str(col).lower()
                    if question_col is None:
                        for pattern in question_patterns:
                            if pattern in col_lower:
                                question_col = col
                                break
                    if answer_col is None:
                        for pattern in answer_patterns:
                            if pattern in col_lower:
                                answer_col = col
                                break

                # Fill in answers
                if question_col and answer_col:
                    for idx, row in df.iterrows():
                        question = str(row[question_col]).strip()
                        if question in answer_map:
                            df.at[idx, answer_col] = answer_map[question]["suggested_answer"]

                # Write sheet
                df.to_excel(writer, sheet_name=sheet_name, index=False)

                # Apply highlighting for low confidence answers
                worksheet = writer.sheets[sheet_name]

                if question_col and answer_col:
                    # Find column indices
                    question_col_idx = list(df.columns).index(question_col) + 1
                    answer_col_idx = list(df.columns).index(answer_col) + 1

                    # Yellow fill for low confidence
                    yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")
                    orange_fill = PatternFill(start_color="FFD580", end_color="FFD580", fill_type="solid")

                    for idx, row in df.iterrows():
                        question = str(row[question_col]).strip()
                        if question in answer_map:
                            qa = answer_map[question]
                            row_num = idx + 2  # +2 for header and 0-indexing

                            if qa["confidence"] < 50:
                                worksheet.cell(row=row_num, column=answer_col_idx).fill = yellow_fill
                            elif qa["confidence"] < 80:
                                worksheet.cell(row=row_num, column=answer_col_idx).fill = orange_fill

        output.seek(0)
        return output.read()

    def export_to_csv(
        self,
        template_content: bytes,
        template_filename: str,
        filled_answers: list[dict],
    ) -> bytes:
        """Export filled answers to CSV format."""
        # Load the original CSV
        df = pd.read_csv(io.BytesIO(template_content))

        # Create a mapping of questions to answers
        answer_map = {qa["question"]: qa for qa in filled_answers}

        # Clean column names
        df.columns = df.columns.str.strip()

        # Find question and answer columns
        question_col = None
        answer_col = None

        question_patterns = ["question", "query", "ask", "q", "requirement", "item"]
        answer_patterns = ["answer", "response", "reply", "a", "value", "input"]

        for col in df.columns:
            col_lower = str(col).lower()
            if question_col is None:
                for pattern in question_patterns:
                    if pattern in col_lower:
                        question_col = col
                        break
            if answer_col is None:
                for pattern in answer_patterns:
                    if pattern in col_lower:
                        answer_col = col
                        break

        # Fill in answers
        if question_col and answer_col:
            for idx, row in df.iterrows():
                question = str(row[question_col]).strip()
                if question in answer_map:
                    df.at[idx, answer_col] = answer_map[question]["suggested_answer"]

        # Convert to CSV
        output = io.BytesIO()
        df.to_csv(output, index=False, encoding='utf-8')
        output.seek(0)
        return output.read()

    def export_to_word(
        self,
        template_content: bytes,
        template_filename: str,
        filled_answers: list[dict],
    ) -> bytes:
        """Export filled answers to Word format."""
        # Load the original Word document
        doc = Document(io.BytesIO(template_content))

        # Create a mapping of questions to answers (with normalized keys)
        answer_map = {}
        for qa in filled_answers:
            # Normalize question text for better matching
            normalized_q = qa["question"].strip().lower()
            answer_map[normalized_q] = qa

        # Try to find and fill in tables
        for table in doc.tables:
            for row in table.rows:
                cells = row.cells
                if len(cells) >= 2:
                    # Get question from first cell
                    question = cells[0].text.strip()
                    normalized_question = question.lower()

                    # Try exact match first, then partial match
                    matched_qa = None
                    if normalized_question in answer_map:
                        matched_qa = answer_map[normalized_question]
                    else:
                        # Try partial matching
                        for key, qa in answer_map.items():
                            if key in normalized_question or normalized_question in key:
                                matched_qa = qa
                                break

                    if matched_qa and matched_qa["suggested_answer"]:
                        # Clear existing content and add new answer
                        cells[1].text = ""  # Clear first
                        paragraph = cells[1].paragraphs[0] if cells[1].paragraphs else cells[1].add_paragraph()
                        run = paragraph.add_run(matched_qa["suggested_answer"])

                        # Highlight based on confidence
                        if matched_qa["confidence"] < 50:
                            run.font.highlight_color = 7  # Yellow
                        elif matched_qa["confidence"] < 80:
                            run.font.color.rgb = RGBColor(255, 165, 0)  # Orange

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output.read()

    def export(
        self,
        template_content: bytes,
        template_filename: str,
        filled_answers: list[dict],
    ) -> tuple[bytes, str]:
        """
        Export filled questionnaire to the same format as the template.

        Returns:
            (file_content, content_type) tuple
        """
        suffix = Path(template_filename).suffix.lower()

        if suffix in [".xlsx", ".xls"]:
            content = self.export_to_excel(template_content, template_filename, filled_answers)
            content_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif suffix == ".csv":
            content = self.export_to_csv(template_content, template_filename, filled_answers)
            content_type = "text/csv"
        elif suffix == ".docx":
            content = self.export_to_word(template_content, template_filename, filled_answers)
            content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            raise ValueError(f"Unsupported export format: {suffix}")

        return content, content_type
